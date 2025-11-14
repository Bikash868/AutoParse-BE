import os
import json
import re

import PyPDF2
import docx
from anthropic import Anthropic


class ResumeParser:
    def __init__(self):
        anthropic_api_key = os.getenv('ANTHROPIC_API_KEY')
        self.anthropic = Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

    def parse_pdf(self, resume_path):
        with open(resume_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            resume_text = ''
            for page in pdf_reader.pages:
                resume_text += page.extract_text()
            return resume_text

    def parse_docx(self, resume_path):
        with open(resume_path, 'rb') as file:
            docx_reader = docx.Document(file)
            resume_text = ''
            for paragraph in docx_reader.paragraphs:
                resume_text += paragraph.text
            return resume_text

    def parse_doc(self, resume_path):
        with open(resume_path, 'rb') as file:
            doc_reader = docx.Document(file)
            resume_text = ''
            for paragraph in doc_reader.paragraphs:
                resume_text += paragraph.text
            return resume_text

    def extract_fields_with_ai(self, resume_text):
        """Using Anthropic API to extract structured fields from resume text."""
        if not self.anthropic:
            # Fallback: return empty dict if no API key
            print('Anthropic API key is not set')
            return {
                'name': None,
                'email': None,
                'phone': None,
                'employer': None,
                'designation': None,
                'skills': None,
                'confidence_scores': None
            }

        prompt = f"""Extract the following information from this resume and return ONLY a valid JSON object with these exact keys:
- name: full name of the candidate
- email: email address
- phone: phone number
- employer: current or most recent employer company name
- designation: current or most recent job title
- skills: comma-separated list of technical and professional skills
- confidence_scores: an object with confidence scores (0-100) for each extracted field

The confidence_scores object should have these keys: name, email, phone, employer, designation, skills
Each confidence score should be a number from 0-100 indicating how confident you are about the extracted information.

If a field is not found, use null for that field and 0 for its confidence score.

Resume text:
{resume_text}

Return only the JSON object, no other text. Example format:
{{
  "name": "John Doe",
  "email": "john@example.com",
  "phone": "+1-234-567-8900",
  "employer": "Tech Corp",
  "designation": "Software Engineer",
  "skills": "Python, Java, React",
  "confidence_scores": {{
    "name": 95,
    "email": 100,
    "phone": 90,
    "employer": 85,
    "designation": 80,
    "skills": 75
  }}
}}"""

        try:
            # print("Calling Anthropic API...")
            message = self.anthropic.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=1024,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            response_text = message.content[0].text.strip()
            print(f"API Response: {response_text}")
            
            # Try to extract JSON from the response
            # Sometimes Claude wraps it in markdown code blocks
            if '```json' in response_text:
                response_text = response_text.split('```json')[1].split('```')[0].strip()
            elif '```' in response_text:
                response_text = response_text.split('```')[1].split('```')[0].strip()
            
            parsed_data = json.loads(response_text)
            # print(f"Parsed data: {parsed_data}")
            return parsed_data
        except Exception as e:
            # Log the error so we can see what went wrong
            print(f"Error in extract_fields_with_ai: {type(e).__name__}: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'name': None,
                'email': None,
                'phone': None,
                'employer': None,
                'designation': None,
                'skills': None,
                'confidence_scores': None
            }

    def parse_resume(self, resume_path):
        if resume_path is None:
            raise ValueError('Resume path is not set')

        if resume_path.endswith('.pdf'):
            resume_text = self.parse_pdf(resume_path)
        elif resume_path.endswith('.docx'):
            resume_text = self.parse_docx(resume_path)
        elif resume_path.endswith('.doc'):
            resume_text = self.parse_doc(resume_path)
        else:
            raise ValueError('Unsupported file type')

        if not resume_text:
            raise ValueError('Failed to parse resume')

        print("resume_text:", resume_text)

        # Extract structured fields using AI
        candidate_data = self.extract_fields_with_ai(resume_text)
        return candidate_data


class AIDocumentRequestGenerator:
    def __init__(self):
        anthropic_api_key = os.getenv('ANTHROPIC_API_KEY')
        self.anthropic = Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

    def generate_request(self, candidate):
        """Generate a personalized document request message for a candidate."""
        if not self.anthropic:
            print('Anthropic API key is not set')
            return self._fallback_request(candidate)
        
        try:
            message = self.anthropic.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=500,
                messages=[{
                    "role": "user",
                    "content": f"""As an AI HR assistant, write a professional and friendly message requesting PAN and Aadhaar documents from a candidate.

Candidate details:
- Name: {candidate.name or 'Candidate'}
- Email: {candidate.email or 'Not provided'}
- Designation: {candidate.designation or 'Not provided'}

Write a personalized message that:
1. Addresses the candidate by name
2. Mentions their application/position if available
3. Politely requests PAN and Aadhaar documents
4. Explains it's for verification purposes
5. Provides clear next steps

Keep it warm, professional, and concise (3-4 sentences)."""
                }]
            )
            
            return message.content[0].text
            
        except Exception as e:
            print(f"AI request generation error: {e}")
            import traceback
            traceback.print_exc()
            return self._fallback_request(candidate)
    
    def _fallback_request(self, candidate):
        """Fallback message when AI is not available."""
        name = candidate.name or 'Candidate'
        designation = candidate.designation or 'your desired position'
        
        return f"""Dear {name},

Thank you for your application for the {designation} position. To proceed with your application, we kindly request you to submit your PAN Card and Aadhaar Card for identity verification.

Please upload the documents at your earliest convenience through our candidate portal.

Best regards,
HR Team"""